from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.core.model import Document, Page, TextElement, ImageElement, TableRegion
from src.core.writer import Writer


class DocxWriter(Writer):
    @staticmethod
    def supported_extensions() -> list[str]:
        return [".docx"]

    def write(self, doc: Document, output_path: str) -> None:
        d = DocxDocument()
        for page in doc.pages:
            self._write_page(d, page)
        d.save(output_path)

    def _write_page(self, d: DocxDocument, page: Page) -> None:
        # page break between pages
        if page.page_number > 0:
            d.add_page_break()

        for element in page.elements:
            if isinstance(element, TextElement):
                self._write_text(d, element)
            elif isinstance(element, ImageElement):
                self._write_image(d, element)
            elif isinstance(element, TableRegion):
                self._write_table(d, element)

    def _write_text(self, d: DocxDocument, te: TextElement) -> None:
        hint = te.hint
        if hint and hint.role == "heading":
            level = min(hint.level, 6) if hint.level >= 1 else 1
            p = d.add_heading(te.text, level=level)
        else:
            p = d.add_paragraph()

        if not (hint and hint.role == "heading"):
            self._apply_text_format(p, te)

    def _apply_text_format(self, p, te: TextElement) -> None:
        run = p.add_run(te.text)
        run.font.size = Pt(te.font_size) if te.font_size else Pt(11)
        run.bold = te.bold
        run.italic = te.italic
        if te.font_name:
            run.font.name = te.font_name
        hex_color = te.font_color.lstrip("#")
        if len(hex_color) == 6:
            run.font.color.rgb = RGBColor(*[int(hex_color[i:i+2], 16) for i in (0, 2, 4)])

    def _write_image(self, d: DocxDocument, ie: ImageElement) -> None:
        import io
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=f".{ie.image_format}", delete=False) as f:
            f.write(ie.image_bytes)
            f.flush()
            # scale to ~4 inches max
            max_w = Inches(4)
            aspect = ie.height / ie.width if ie.width else 1
            w = min(max_w, Inches(ie.width / 72)) if ie.width else max_w
            h = w * aspect
            d.add_picture(f.name, width=w, height=h)

    def _write_table(self, d: DocxDocument, tr: TableRegion) -> None:
        if not tr.rows:
            return
        rows = len(tr.rows)
        cols = max(len(row) for row in tr.rows) if tr.rows else 1
        table = d.add_table(rows=rows, cols=cols, style="Table Grid")
        for i, row in enumerate(tr.rows):
            for j, cell_text in enumerate(row):
                if j < cols:
                    table.cell(i, j).text = cell_text
